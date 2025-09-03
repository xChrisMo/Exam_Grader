"""
Consolidated File Processing Service

This service consolidates all file processing functionality from multiple services.
"""

import hashlib
import os
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from src.constants import DEFAULT_MAX_FILE_SIZE_MB, DIR_TEMP, SUPPORTED_EXTENSIONS
from src.services.base_service import BaseService, ServiceStatus
from utils.logger import logger


class FileProcessingService(BaseService):
    """Consolidated file processing service"""
    
    def __init__(self, max_file_size_mb: int = DEFAULT_MAX_FILE_SIZE_MB):
        super().__init__("file_processing_service")
        self.max_file_size_mb = max_file_size_mb
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        
        # Processing statistics
        self.processing_stats = {
            'total_files': 0,
            'successful_extractions': 0,
            'failed_extractions': 0,
        }
        
    async def initialize(self) -> bool:
        """Initialize the file processing service"""
        try:
            temp_dir = Path(DIR_TEMP)
            temp_dir.mkdir(exist_ok=True)
            
            self.status = ServiceStatus.HEALTHY
            logger.info("File processing service initialized successfully")
            return True
        except Exception as e:
            self.status = ServiceStatus.UNHEALTHY
            logger.error(f"Failed to initialize file processing service: {str(e)}")
            return False

    async def health_check(self) -> bool:
        """Perform health check"""
        try:
            temp_dir = Path(DIR_TEMP)
            return temp_dir.exists() and os.access(temp_dir, os.W_OK)
        except Exception as e:
            logger.error(f"File processing service health check failed: {str(e)}")
            return False

    async def cleanup(self) -> None:
        """Clean up resources"""
        try:
            logger.info("File processing service cleanup completed")
        except Exception as e:
            logger.error(f"Error during file processing service cleanup: {str(e)}")
    
    def process_file_with_fallback(self, file_path: str, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Process file with fallback methods"""
        start_time = time.time()
        
        try:
            self.processing_stats['total_files'] += 1
            
            # Validate file
            if not self._validate_file(file_path):
                return self._create_error_result(file_path, start_time, "File validation failed")
            
            # Extract content
            content = self._extract_content(file_path)
            
            if content:
                self.processing_stats['successful_extractions'] += 1
                return self._create_success_result(file_path, start_time, content)
            else:
                self.processing_stats['failed_extractions'] += 1
                return self._create_error_result(file_path, start_time, "No content extracted")
            
        except Exception as e:
            self.processing_stats['failed_extractions'] += 1
            return self._create_error_result(file_path, start_time, str(e))
    
    def _validate_file(self, file_path: str) -> bool:
        """Validate file before processing"""
        try:
            if not os.path.exists(file_path):
                return False
            
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size_bytes:
                return False
            
            file_extension = Path(file_path).suffix.lower()
            return file_extension in SUPPORTED_EXTENSIONS
            
        except Exception:
            return False
    
    def _extract_content(self, file_path: str) -> str:
        """Extract content from file with proper document processing"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.txt':
                return self._extract_text(file_path)
            elif file_extension == '.pdf':
                return self._extract_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_docx(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif']:
                return self._extract_image_ocr(file_path)
            else:
                # Fallback to text extraction
                return self._extract_text(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return ""
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except:
                    continue
            return ""
        except Exception:
            return ""
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF files using OCR"""
        try:
            from src.parsing.parse_submission import DocumentParser
            
            # Use the OCR-based PDF extraction
            text = DocumentParser.extract_text_from_pdf(file_path)
            
            if text and text.strip():
                return text
            else:
                logger.warning("No text extracted from PDF")
                return ""
            
        except Exception as e:
            logger.error(f"Error extracting PDF content: {e}")
            return ""
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from Word documents"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n\n'.join(text_content)
            
        except ImportError:
            logger.warning("python-docx not available, cannot extract Word document content")
            return ""
        except Exception as e:
            logger.error(f"Error extracting Word document content: {e}")
            return ""
    
    def _extract_image_ocr(self, file_path: str) -> str:
        """Extract text from images using HandwritingOCR"""
        try:
            from src.parsing.parse_submission import DocumentParser
            
            # Use the HandwritingOCR-based image extraction
            text = DocumentParser.extract_text_from_image(file_path)
            
            if text and text.strip():
                return text
            else:
                logger.warning("No text extracted from image")
                return ""
            
        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            return ""
    
    def _create_success_result(self, file_path: str, start_time: float, content: str) -> Dict[str, Any]:
        """Create successful processing result"""
        return {
            'success': True,
            'text_content': content,
            'word_count': len(content.split()),
            'character_count': len(content),
            'processing_duration_ms': int((time.time() - start_time) * 1000),
            'processing_timestamp': datetime.now(timezone.utc).isoformat()
        }
    
    def _create_error_result(self, file_path: str, start_time: float, error_message: str) -> Dict[str, Any]:
        """Create error processing result"""
        return {
            'success': False,
            'text_content': '',
            'word_count': 0,
            'character_count': 0,
            'processing_duration_ms': int((time.time() - start_time) * 1000),
            'error_message': error_message,
            'processing_timestamp': datetime.now(timezone.utc).isoformat()
        }
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Get processing statistics"""
        stats = self.processing_stats.copy()
        stats['success_rate'] = (
            stats['successful_extractions'] / max(1, stats['total_files']) * 100
        )
        return stats


# Global instance
file_processing_service = FileProcessingService()