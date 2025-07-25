"""Document processing service for LLM training system."""
from typing import Any, Dict, List, Optional

import json
import uuid
import logging
from io import BytesIO
import re

# Document processing libraries
try:
    import PyMuPDF as fitz  # For PDF processing
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    fitz = None

from docx import Document as DocxDocument  # For DOCX processing

from ..services.base_service import BaseService
from ..models.document_models import (
    DocumentType, DocumentStatus, ProcessedDocument, DocumentMetadata,
    DocumentProcessingResult, FileUpload, Dataset
)

# Define ValidationResult locally to avoid circular imports
class ValidationResult:
    """Validation result helper class"""
    
    def __init__(self, is_valid: bool = True):
        self.errors = []
        self.warnings = []
        self._is_valid = is_valid
    
    @property
    def is_valid(self) -> bool:
        return len(self.errors) == 0 and self._is_valid
    
    def add_error(self, field: str, message: str, code: str, value: Any = None):
        self.errors.append({
            'field': field,
            'message': message,
            'code': code,
            'value': value
        })
        self._is_valid = False
    
    def add_warning(self, field: str, message: str, suggestion: str = None):
        self.warnings.append({
            'field': field,
            'message': message,
            'suggestion': suggestion
        })
    
    def get_error_summary(self) -> str:
        if not self.errors:
            return ""
        return "; ".join([error['message'] for error in self.errors])

logger = logging.getLogger(__name__)


class DocumentProcessorService(BaseService):
    """Service for processing and managing training documents."""
    
    # Configuration constants
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    SUPPORTED_EXTENSIONS = ['pdf', 'txt', 'docx', 'json']
    MAX_CONTENT_LENGTH = 10 * 1024 * 1024  # 10MB of text content
    
    def __init__(self, **kwargs):
        """Initialize document processor service."""
        super().__init__("document_processor", **kwargs)
        self._documents: Dict[str, ProcessedDocument] = {}
        self._datasets: Dict[str, Dataset] = {}
        
    def initialize(self) -> bool:
        """Initialize the document processor service."""
        try:
            logger.info("Initializing DocumentProcessorService")
            # Test document processing libraries
            self._test_dependencies()
            logger.info("DocumentProcessorService initialized successfully")
            return True
        except Exception as e:
            logger.error(f"Failed to initialize DocumentProcessorService: {str(e)}")
            return False
    
    def health_check(self) -> bool:
        """Perform health check for document processor."""
        try:
            # Test basic functionality
            test_content = "Test document content"
            result = self._sanitize_content(test_content)
            return result is not None
        except Exception as e:
            logger.error(f"Document processor health check failed: {str(e)}")
            return False
    
    def cleanup(self) -> None:
        """Cleanup document processor resources."""
        self._documents.clear()
        self._datasets.clear()
        logger.info("DocumentProcessorService cleaned up")
    
    def _test_dependencies(self) -> None:
        """Test that all required dependencies are available."""
        try:
            # Test PyMuPDF
            if PYMUPDF_AVAILABLE:
                fitz.Document()
                logger.info("PyMuPDF available for PDF processing")
            else:
                logger.warning("PyMuPDF not available - PDF processing disabled")
            
            # Test python-docx
            DocxDocument()
            
            logger.info("Document processing dependencies checked")
        except Exception as e:
            raise RuntimeError(f"Document processing dependencies not available: {str(e)}")
    
    def validate_file_upload(self, file_upload: FileUpload) -> ValidationResult:
        """Validate file upload before processing.
        
        Args:
            file_upload: File upload information
            
        Returns:
            ValidationResult with any validation errors
        """
        result = ValidationResult(is_valid=True)
        
        # Check file size
        if file_upload.size > self.MAX_FILE_SIZE:
            result.add_error(
                field="file_size",
                message=f"File size ({file_upload.size / (1024*1024):.1f}MB) exceeds maximum allowed size ({self.MAX_FILE_SIZE / (1024*1024):.1f}MB)",
                code="FILE_TOO_LARGE"
            )
        
        # Check file extension
        extension = file_upload.get_file_extension()
        if extension not in self.SUPPORTED_EXTENSIONS:
            result.add_error(
                field="file_extension",
                message=f"File extension '{extension}' is not supported. Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}",
                code="UNSUPPORTED_FILE_TYPE"
            )
        
        # Check filename
        if not file_upload.filename or not file_upload.filename.strip():
            result.add_error(
                field="filename",
                message="Filename is required",
                code="MISSING_FILENAME"
            )
        
        # Check for potentially dangerous filenames
        if self._is_dangerous_filename(file_upload.filename):
            result.add_error(
                field="filename",
                message="Filename contains potentially dangerous characters",
                code="DANGEROUS_FILENAME"
            )
        
        return result
    
    def process_file_upload(self, file_upload: FileUpload) -> DocumentProcessingResult:
        """Process uploaded file and extract text content.
        
        Args:
            file_upload: File upload information
            
        Returns:
            DocumentProcessingResult with processed document or error
        """
        with self.track_request():
            try:
                # Validate file upload
                validation_result = self.validate_file_upload(file_upload)
                if not validation_result.is_valid:
                    error_messages = validation_result.get_error_messages()
                    return DocumentProcessingResult.error_result(
                        error_message="; ".join(error_messages)
                    )
                
                # Get document type
                document_type = file_upload.get_document_type()
                if not document_type:
                    return DocumentProcessingResult.error_result(
                        error_message=f"Unsupported file type: {file_upload.get_file_extension()}"
                    )
                
                # Extract text content based on file type
                content = self._extract_text_content(file_upload, document_type)
                
                # Sanitize content
                sanitized_content = self._sanitize_content(content)
                
                # Check content length
                if len(sanitized_content) > self.MAX_CONTENT_LENGTH:
                    return DocumentProcessingResult.error_result(
                        error_message=f"Extracted content too large ({len(sanitized_content)} chars). Maximum allowed: {self.MAX_CONTENT_LENGTH} chars"
                    )
                
                # Create processed document
                document_id = str(uuid.uuid4())
                processed_doc = ProcessedDocument.create_from_file(
                    file_id=document_id,
                    filename=file_upload.filename,
                    content=sanitized_content,
                    document_type=document_type,
                    file_size=file_upload.size
                )
                
                # Store document
                self._documents[document_id] = processed_doc
                
                logger.info(f"Successfully processed document: {file_upload.filename} ({document_type.value})")
                
                return DocumentProcessingResult.success_result(processed_doc)
                
            except Exception as e:
                logger.error(f"Error processing file upload {file_upload.filename}: {str(e)}")
                return DocumentProcessingResult.error_result(
                    error_message=f"Failed to process document: {str(e)}"
                )
    
    def _extract_text_content(self, file_upload: FileUpload, document_type: DocumentType) -> str:
        """Extract text content from uploaded file based on document type.
        
        Args:
            file_upload: File upload information
            document_type: Type of document to process
            
        Returns:
            Extracted text content
        """
        if document_type == DocumentType.PDF:
            return self._extract_pdf_content(file_upload.content)
        elif document_type == DocumentType.TXT:
            return self._extract_txt_content(file_upload.content)
        elif document_type == DocumentType.DOCX:
            return self._extract_docx_content(file_upload.content)
        elif document_type == DocumentType.JSON:
            return self._extract_json_content(file_upload.content)
        else:
            raise ValueError(f"Unsupported document type: {document_type}")
    
    def _extract_pdf_content(self, content: bytes) -> str:
        """Extract text from PDF content.
        
        Args:
            content: PDF file content as bytes
            
        Returns:
            Extracted text content
        """
        try:
            if not PYMUPDF_AVAILABLE:
                raise RuntimeError("PDF processing not available. Install PyMuPDF: pip install PyMuPDF")
                
            # Open PDF from bytes
            pdf_document = fitz.open(stream=content, filetype="pdf")
            
            text_content = []
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text_content.append(page.get_text())
            
            pdf_document.close()
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise RuntimeError(f"Failed to extract PDF content: {str(e)}")
    
    def _extract_txt_content(self, content: bytes) -> str:
        """Extract text from TXT content.
        
        Args:
            content: TXT file content as bytes
            
        Returns:
            Extracted text content
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            return content.decode('utf-8', errors='replace')
            
        except Exception as e:
            raise RuntimeError(f"Failed to extract TXT content: {str(e)}")
    
    def _extract_docx_content(self, content: bytes) -> str:
        """Extract text from DOCX content.
        
        Args:
            content: DOCX file content as bytes
            
        Returns:
            Extracted text content
        """
        try:
            # Open DOCX from bytes
            doc = DocxDocument(BytesIO(content))
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise RuntimeError(f"Failed to extract DOCX content: {str(e)}")
    
    def _extract_json_content(self, content: bytes) -> str:
        """Extract text from JSON content.
        
        Args:
            content: JSON file content as bytes
            
        Returns:
            Extracted text content (formatted JSON or specific fields)
        """
        try:
            # Decode JSON content
            json_str = content.decode('utf-8')
            json_data = json.loads(json_str)
            
            # If it's a simple structure, extract text values
            if isinstance(json_data, dict):
                text_content = self._extract_text_from_dict(json_data)
            elif isinstance(json_data, list):
                text_content = self._extract_text_from_list(json_data)
            else:
                # For simple values, convert to string
                text_content = str(json_data)
            
            return text_content
            
        except json.JSONDecodeError as e:
            raise RuntimeError(f"Invalid JSON format: {str(e)}")
        except Exception as e:
            raise RuntimeError(f"Failed to extract JSON content: {str(e)}")
    
    def _extract_text_from_dict(self, data: dict, prefix: str = "") -> str:
        """Extract text content from dictionary recursively."""
        text_parts = []
        
        for key, value in data.items():
            current_key = f"{prefix}.{key}" if prefix else key
            
            if isinstance(value, str):
                text_parts.append(f"{current_key}: {value}")
            elif isinstance(value, dict):
                text_parts.append(self._extract_text_from_dict(value, current_key))
            elif isinstance(value, list):
                text_parts.append(self._extract_text_from_list(value, current_key))
            else:
                text_parts.append(f"{current_key}: {str(value)}")
        
        return "\n".join(filter(None, text_parts))
    
    def _extract_text_from_list(self, data: list, prefix: str = "") -> str:
        """Extract text content from list recursively."""
        text_parts = []
        
        for i, item in enumerate(data):
            current_key = f"{prefix}[{i}]" if prefix else f"item_{i}"
            
            if isinstance(item, str):
                text_parts.append(f"{current_key}: {item}")
            elif isinstance(item, dict):
                text_parts.append(self._extract_text_from_dict(item, current_key))
            elif isinstance(item, list):
                text_parts.append(self._extract_text_from_list(item, current_key))
            else:
                text_parts.append(f"{current_key}: {str(item)}")
        
        return "\n".join(filter(None, text_parts))
    
    def _sanitize_content(self, content: str) -> str:
        """Sanitize extracted content for training.
        
        Args:
            content: Raw extracted content
            
        Returns:
            Sanitized content
        """
        if not content:
            return ""
        
        # Remove excessive whitespace
        content = re.sub(r'\s+', ' ', content)
        
        # Remove control characters except newlines and tabs
        content = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', content)
        
        # Normalize line endings
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive blank lines
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        
        # Strip leading/trailing whitespace
        content = content.strip()
        
        return content
    
    def _is_dangerous_filename(self, filename: str) -> bool:
        """Check if filename contains potentially dangerous characters.
        
        Args:
            filename: Filename to check
            
        Returns:
            True if filename is potentially dangerous
        """
        dangerous_patterns = [
            r'\.\./',  # Directory traversal
            r'[<>:"|?*]',  # Windows reserved characters
            r'^(CON|PRN|AUX|NUL|COM[1-9]|LPT[1-9])(\.|$)',  # Windows reserved names
        ]
        
        for pattern in dangerous_patterns:
            if re.search(pattern, filename, re.IGNORECASE):
                return True
        
        return False

    def get_document(self, document_id: str) -> Optional[ProcessedDocument]:
        """Get a processed document by ID.
        
        Args:
            document_id: Document ID
            
        Returns:
            ProcessedDocument or None if not found
        """
        return self._documents.get(document_id)
    
    def list_documents(self) -> List[ProcessedDocument]:
        """Get list of all processed documents.
        
        Returns:
            List of ProcessedDocument objects
        """
        return list(self._documents.values())
    
    def delete_document(self, document_id: str) -> bool:
        """Delete a processed document.
        
        Args:
            document_id: Document ID to delete
            
        Returns:
            True if document was deleted, False if not found
        """
        if document_id in self._documents:
            # Remove from all datasets
            for dataset in self._datasets.values():
                dataset.remove_document(document_id)
            
            # Remove document
            del self._documents[document_id]
            logger.info(f"Deleted document: {document_id}")
            return True
        
        return False
    
    # Dataset management methods
    
    def create_dataset(self, name: str, description: str = None, document_ids: List[str] = None) -> Dataset:
        """Create a new document dataset.
        
        Args:
            name: Dataset name
            description: Optional dataset description
            document_ids: Optional list of document IDs to include
            
        Returns:
            Created Dataset object
        """
        dataset_id = str(uuid.uuid4())
        dataset = Dataset(
            id=dataset_id,
            name=name,
            description=description,
            document_ids=document_ids or []
        )
        
        # Validate document IDs exist
        valid_doc_ids = []
        for doc_id in dataset.document_ids:
            if doc_id in self._documents:
                valid_doc_ids.append(doc_id)
                # Add dataset reference to document
                self._documents[doc_id].datasets.append(dataset_id)
        
        dataset.document_ids = valid_doc_ids
        self._datasets[dataset_id] = dataset
        
        logger.info(f"Created dataset: {name} with {len(valid_doc_ids)} documents")
        return dataset
    
    def get_dataset(self, dataset_id: str) -> Optional[Dataset]:
        """Get a dataset by ID.
        
        Args:
            dataset_id: Dataset ID
            
        Returns:
            Dataset or None if not found
        """
        return self._datasets.get(dataset_id)
    
    def list_datasets(self) -> List[Dataset]:
        """Get list of all datasets.
        
        Returns:
            List of Dataset objects
        """
        return list(self._datasets.values())
    
    def add_document_to_dataset(self, dataset_id: str, document_id: str) -> bool:
        """Add a document to a dataset.
        
        Args:
            dataset_id: Dataset ID
            document_id: Document ID to add
            
        Returns:
            True if document was added, False if dataset/document not found
        """
        if dataset_id not in self._datasets or document_id not in self._documents:
            return False
        
        dataset = self._datasets[dataset_id]
        document = self._documents[document_id]
        
        dataset.add_document(document_id)
        if dataset_id not in document.datasets:
            document.datasets.append(dataset_id)
        
        return True
    
    def remove_document_from_dataset(self, dataset_id: str, document_id: str) -> bool:
        """Remove a document from a dataset.
        
        Args:
            dataset_id: Dataset ID
            document_id: Document ID to remove
            
        Returns:
            True if document was removed, False if dataset/document not found
        """
        if dataset_id not in self._datasets or document_id not in self._documents:
            return False
        
        dataset = self._datasets[dataset_id]
        document = self._documents[document_id]
        
        dataset.remove_document(document_id)
        if dataset_id in document.datasets:
            document.datasets.remove(dataset_id)
        
        return True
    
    def delete_dataset(self, dataset_id: str) -> bool:
        """Delete a dataset.
        
        Args:
            dataset_id: Dataset ID to delete
            
        Returns:
            True if dataset was deleted, False if not found
        """
        if dataset_id not in self._datasets:
            return False
        
        dataset = self._datasets[dataset_id]
        
        # Remove dataset reference from all documents
        for doc_id in dataset.document_ids:
            if doc_id in self._documents:
                document = self._documents[doc_id]
                if dataset_id in document.datasets:
                    document.datasets.remove(dataset_id)
        
        # Delete dataset
        del self._datasets[dataset_id]
        logger.info(f"Deleted dataset: {dataset_id}")
        return True
    
    def get_dataset_statistics(self, dataset_id: str) -> Optional[Dict[str, Any]]:
        """Get statistics for a dataset.
        
        Args:
            dataset_id: Dataset ID
            
        Returns:
            Dictionary with dataset statistics or None if not found
        """
        if dataset_id not in self._datasets:
            return None
        
        dataset = self._datasets[dataset_id]
        documents = [self._documents[doc_id] for doc_id in dataset.document_ids 
                    if doc_id in self._documents]
        
        total_words = sum(doc.metadata.word_count for doc in documents)
        total_chars = sum(doc.metadata.character_count for doc in documents)
        total_size = sum(doc.metadata.file_size for doc in documents)
        
        doc_types = {}
        for doc in documents:
            doc_type = doc.document_type.value
            doc_types[doc_type] = doc_types.get(doc_type, 0) + 1
        
        return {
            "dataset_id": dataset_id,
            "name": dataset.name,
            "document_count": len(documents),
            "total_words": total_words,
            "total_characters": total_chars,
            "total_file_size": total_size,
            "document_types": doc_types,
            "created_date": dataset.created_date.isoformat(),
            "updated_date": dataset.updated_date.isoformat()
        }