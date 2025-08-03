"""
LLM Training Routes

This module provides routes for the LLM training and fine-tuning interface.
"""

from flask import Blueprint, render_template, request, jsonify, current_app, send_file
from flask_login import login_required, current_user
from src.database.models import db, LLMDataset, LLMTrainingJob, LLMTrainingReport, LLMDocument, LLMDatasetDocument
from src.services.llm_training_service import LLMTrainingService
from utils.logger import logger
import traceback
import os
import re
import uuid
from datetime import datetime, timezone

def extract_text_from_file(file_path, file_extension):
    """
    Enhanced text extraction with comprehensive error handling and fallback mechanisms.
    
    Args:
        file_path: Path to the file to extract text from
        file_extension: File extension (e.g., '.pdf', '.docx')
        
    Returns:
        Extracted text content as string
    """
    from src.services.file_processing_service import FileProcessingService
    from src.services.error_handling_service import LLMTrainingErrorHandler
    
    try:
        # Use the enhanced file processing service
        file_processor = FileProcessingService()
        
        file_info = {
            'name': os.path.basename(file_path),
            'path': file_path,
            'extension': file_extension,
            'size': os.path.getsize(file_path) if os.path.exists(file_path) else 0
        }
        
        # Process file with fallback mechanisms
        result = file_processor.process_file_with_fallback(file_path, file_info)
        
        if result['success'] and result.get('text_content', '').strip():
            logger.info(f"Successfully extracted {result['word_count']} words from {file_info['name']} using {result['extraction_method']} method")
            return result['text_content']
        else:
            # Log extraction failure with details
            error_msg = f"Failed to extract text from {file_info['name']}: {'; '.join(result.get('validation_errors', ['Unknown error']))}"
            logger.warning(error_msg)
            
            # Try OCR extraction for supported formats
            if file_extension.lower() in ['.pdf', '.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif']:
                try:
                    logger.info(f"Attempting OCR extraction for {file_info['name']}")
                    from src.services.consolidated_ocr_service import ConsolidatedOCRService
                    ocr_service = ConsolidatedOCRService()
                    ocr_result = ocr_service.extract_text(file_path)
                    
                    if ocr_result and ocr_result.get('text', '').strip():
                        logger.info(f"Successfully extracted text using OCR from {file_info['name']}")
                        return ocr_result['text']
                    else:
                        logger.warning(f"OCR extraction returned empty text for {file_info['name']}")
                        
                except Exception as ocr_error:
                    logger.error(f"OCR extraction failed for {file_info['name']}: {ocr_error}")
            
            # Try legacy extraction as final fallback
            return _legacy_extract_text_from_file(file_path, file_extension)
            
    except Exception as e:
        logger.error(f"Error in enhanced text extraction for {file_path}: {e}")
        
        # Try OCR as fallback for supported formats before legacy extraction
        if file_extension.lower() in ['.pdf', '.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif']:
            try:
                logger.info(f"Attempting OCR fallback for {file_path}")
                from src.services.consolidated_ocr_service import ConsolidatedOCRService
                ocr_service = ConsolidatedOCRService()
                ocr_result = ocr_service.extract_text(file_path)
                
                if ocr_result and ocr_result.get('text', '').strip():
                    logger.info(f"OCR fallback successful for {file_path}")
                    return ocr_result['text']
                    
            except Exception as ocr_error:
                logger.error(f"OCR fallback failed for {file_path}: {ocr_error}")
        
        # Final fallback to legacy extraction
        return _legacy_extract_text_from_file(file_path, file_extension)

def _legacy_extract_text_from_file(file_path, file_extension):
    """
    Legacy text extraction function as final fallback.
    Kept for compatibility and as last resort extraction method.
    """
    text_content = ""
    
    try:
        if file_extension == '.txt':
            # Plain text files with encoding detection
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text_content = f.read()
                        break
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    logger.debug(f"Failed to read with encoding {encoding}: {e}")
                    continue
                    
        elif file_extension == '.docx':
            # Word documents with multiple fallback methods
            try:
                import docx
                doc = docx.Document(file_path)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += '\n' + cell.text
                            
            except ImportError:
                logger.warning("python-docx not installed, trying alternative extraction")
                try:
                    # Try docx2txt as alternative
                    import docx2txt
                    text_content = docx2txt.process(file_path)
                except ImportError:
                    logger.warning("docx2txt not available, using basic text extraction")
                    try:
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            text_content = f.read()
                    except:
                        text_content = ""
            except Exception as e:
                logger.warning(f"Error extracting from DOCX: {e}")
                text_content = ""
                    
        elif file_extension == '.pdf':
            # PDF files with multiple extraction methods
            extraction_methods = [
                ('PyPDF2', _extract_pdf_pypdf2),
                ('pdfplumber', _extract_pdf_pdfplumber),
                ('pdfminer', _extract_pdf_pdfminer)
            ]
            
            for method_name, method_func in extraction_methods:
                try:
                    text_content = method_func(file_path)
                    if text_content and text_content.strip():
                        logger.info(f"Successfully extracted PDF using {method_name}")
                        break
                except ImportError:
                    logger.debug(f"{method_name} not available")
                    continue
                except Exception as e:
                    logger.debug(f"PDF extraction with {method_name} failed: {e}")
                    continue
                    
        elif file_extension in ['.doc']:
            # Legacy Word documents
            extraction_methods = [
                ('antiword', _extract_doc_antiword),
                ('catdoc', _extract_doc_catdoc),
                ('textract', _extract_doc_textract)
            ]
            
            for method_name, method_func in extraction_methods:
                try:
                    text_content = method_func(file_path)
                    if text_content and text_content.strip():
                        logger.info(f"Successfully extracted DOC using {method_name}")
                        break
                except Exception as e:
                    logger.debug(f"DOC extraction with {method_name} failed: {e}")
                    continue
                    
        elif file_extension in ['.rtf']:
            # RTF files with fallback
            try:
                from striprtf.striprtf import rtf_to_text
                with open(file_path, 'r', encoding='utf-8') as f:
                    rtf_content = f.read()
                text_content = rtf_to_text(rtf_content)
            except ImportError:
                logger.warning("striprtf not installed, using regex-based RTF extraction")
                text_content = _extract_rtf_regex(file_path)
            except Exception as e:
                logger.warning(f"Error extracting from RTF: {e}")
                text_content = _extract_rtf_regex(file_path)
                
        elif file_extension in ['.html', '.htm']:
            # HTML files with BeautifulSoup and fallback
            try:
                from bs4 import BeautifulSoup
                with open(file_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                    
                text_content = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text_content.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text_content = '\n'.join(chunk for chunk in chunks if chunk)
                
            except ImportError:
                logger.warning("BeautifulSoup not installed, using regex-based HTML extraction")
                text_content = _extract_html_regex(file_path)
            except Exception as e:
                logger.warning(f"Error extracting from HTML: {e}")
                text_content = _extract_html_regex(file_path)
                
        elif file_extension in ['.md', '.markdown']:
            # Markdown files with enhanced processing
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
                
                # Enhanced markdown cleanup
                text_content = _clean_markdown_content(text_content)
                
            except Exception as e:
                logger.warning(f"Error extracting from Markdown: {e}")
                text_content = ""
                
        elif file_extension in ['.json']:
            # JSON files with structured text extraction
            try:
                text_content = _extract_json_text(file_path)
            except Exception as e:
                logger.warning(f"Error extracting from JSON: {e}")
                text_content = ""
                
        else:
            # Unknown file type - comprehensive fallback
            text_content = _extract_unknown_format(file_path, file_extension)
    
    except Exception as e:
        logger.error(f"Error in legacy text extraction from {file_path}: {e}")
        text_content = ""
    
    # Enhanced text cleanup
    if text_content:
        text_content = _clean_extracted_text(text_content)
    
    return text_content

def _extract_pdf_pypdf2(file_path):
    """Extract text from PDF using PyPDF2"""
    import PyPDF2
    with open(file_path, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return '\n'.join(text_parts)

def _extract_pdf_pdfplumber(file_path):
    """Extract text from PDF using pdfplumber"""
    import pdfplumber
    with pdfplumber.open(file_path) as pdf:
        text_parts = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return '\n'.join(text_parts)

def _extract_pdf_pdfminer(file_path):
    """Extract text from PDF using pdfminer"""
    from pdfminer.high_level import extract_text
    return extract_text(file_path)

def _extract_doc_antiword(file_path):
    """Extract text from DOC using antiword"""
    import subprocess
    result = subprocess.run(['antiword', file_path], 
                          capture_output=True, text=True, timeout=30)
    if result.returncode == 0:
        return result.stdout
    else:
        raise ValueError("antiword extraction failed")

def _extract_doc_catdoc(file_path):
    """Extract text from DOC using catdoc"""
    import subprocess
    result = subprocess.run(['catdoc', file_path], 
                          capture_output=True, text=True, timeout=30)
    if result.returncode == 0:
        return result.stdout
    else:
        raise ValueError("catdoc extraction failed")

def _extract_doc_textract(file_path):
    """Extract text from DOC using textract"""
    import textract
    return textract.process(file_path).decode('utf-8')

def _extract_rtf_regex(file_path):
    """Extract text from RTF using regex-based approach"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    # Remove RTF control codes and formatting
    content = re.sub(r'\\[a-z]+\d*\s?', '', content)  # Remove RTF commands
    content = re.sub(r'[{}]', '', content)  # Remove braces
    content = re.sub(r'\s+', ' ', content)  # Normalize whitespace
    return content.strip()

def _extract_html_regex(file_path):
    """Extract text from HTML using regex-based approach"""
    with open(file_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Remove script and style content
    html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
    html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove HTML tags
    text_content = re.sub(r'<[^>]+>', '', html_content)
    
    # Decode HTML entities
    import html
    text_content = html.unescape(text_content)
    
    return text_content

def _clean_markdown_content(content):
    """Enhanced markdown content cleaning"""
    # Remove headers
    content = re.sub(r'^#{1,6}\s+(.*)$', r'\1', content, flags=re.MULTILINE)
    
    # Remove bold and italic formatting
    content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Bold
    content = re.sub(r'\*(.*?)\*', r'\1', content)  # Italic
    content = re.sub(r'__(.*?)__', r'\1', content)  # Bold alternative
    content = re.sub(r'_(.*?)_', r'\1', content)  # Italic alternative
    
    # Remove code blocks and inline code
    content = re.sub(r'```.*?```', '', content, flags=re.DOTALL)  # Code blocks
    content = re.sub(r'`(.*?)`', r'\1', content)  # Inline code
    
    # Remove links but keep text
    content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)  # Links
    content = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', content)  # Images
    
    # Remove horizontal rules
    content = re.sub(r'^[-*_]{3,}$', '', content, flags=re.MULTILINE)
    
    # Clean up lists
    content = re.sub(r'^\s*[-*+]\s+', '', content, flags=re.MULTILINE)  # Unordered lists
    content = re.sub(r'^\s*\d+\.\s+', '', content, flags=re.MULTILINE)  # Ordered lists
    
    return content

def _extract_json_text(file_path):
    """Extract text content from JSON files"""
    import json
    
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    def extract_text_values(obj, texts=None):
        if texts is None:
            texts = []
            
        if isinstance(obj, dict):
            for value in obj.values():
                extract_text_values(value, texts)
        elif isinstance(obj, list):
            for item in obj:
                extract_text_values(item, texts)
        elif isinstance(obj, str) and len(obj.strip()) > 0:
            texts.append(obj.strip())
        elif obj is not None:
            texts.append(str(obj))
            
        return texts
    
    text_values = extract_text_values(data)
    return '\n'.join(text_values)

def _extract_unknown_format(file_path, file_extension):
    """Extract text from unknown file formats using multiple approaches"""
    text_content = ""
    
    # Try different encoding approaches
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'ascii']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read()
                
            if _is_likely_text_content(content):
                text_content = content
                logger.info(f"Successfully read unknown format {file_extension} with encoding {encoding}")
                break
                
        except Exception as e:
            logger.debug(f"Failed to read with encoding {encoding}: {e}")
            continue
    
    # If no encoding worked, try binary approach
    if not text_content:
        try:
            with open(file_path, 'rb') as f:
                binary_content = f.read()
                
            # Try to decode as text, ignoring errors
            text_content = binary_content.decode('utf-8', errors='ignore')
            
            # Filter out non-printable characters
            text_content = ''.join(char for char in text_content if char.isprintable() or char.isspace())
            
        except Exception as e:
            logger.warning(f"Binary extraction failed for {file_path}: {e}")
    
    return text_content

def _is_likely_text_content(content):
    """Check if content is likely to be readable text"""
    if not content:
        return False
    
    # Check ratio of printable characters
    printable_chars = sum(1 for char in content if char.isprintable() or char.isspace())
    printable_ratio = printable_chars / len(content)
    
    has_words = bool(re.search(r'\b[a-zA-Z]{3,}\b', content))
    has_sentences = bool(re.search(r'[.!?]\s+[A-Z]', content))
    
    return printable_ratio > 0.7 and (has_words or has_sentences)

def _clean_extracted_text(text_content):
    """Enhanced text cleaning with better formatting preservation"""
    if not text_content:
        return ""
    
    # Remove control characters but preserve newlines and tabs
    text_content = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text_content)
    
    # Normalize different types of whitespace
    text_content = re.sub(r'[\u00A0\u1680\u2000-\u200B\u202F\u205F\u3000]', ' ', text_content)
    
    text_content = re.sub(r'[ \t]+', ' ', text_content)  # Multiple spaces/tabs to single space
    text_content = re.sub(r'\n[ \t]+', '\n', text_content)  # Remove leading whitespace on lines
    text_content = re.sub(r'[ \t]+\n', '\n', text_content)  # Remove trailing whitespace on lines
    text_content = re.sub(r'\n{3,}', '\n\n', text_content)  # Multiple newlines to double newline
    
    # Remove common extraction artifacts
    text_content = re.sub(r'\f', '\n', text_content)  # Form feeds to newlines
    text_content = re.sub(r'[\u200B-\u200D\uFEFF]', '', text_content)  # Zero-width characters
    
    return text_content.strip()

# Create blueprint
llm_training_bp = Blueprint('llm_training', __name__, url_prefix='/llm-training')

@llm_training_bp.route('/')
@login_required
def index():
    """Main LLM training page"""
    try:
        return render_template('llm_training.html')
    except Exception as e:
        logger.error(f"Error rendering LLM training page: {e}")
        logger.error(traceback.format_exc())
        return render_template('error.html', 
                             error_message="Failed to load LLM training page"), 500

@llm_training_bp.route('/api/config')
@login_required
def get_config():
    """Get LLM training configuration"""
    try:
        config = {
            'models': [
                {'id': 'gpt-3.5-turbo', 'name': 'GPT-3.5 Turbo', 'provider': 'openai'},
                {'id': 'gpt-4', 'name': 'GPT-4', 'provider': 'openai'},
                {'id': 'deepseek-chat', 'name': 'DeepSeek Chat', 'provider': 'deepseek'},
                {'id': 'deepseek-reasoner', 'name': 'DeepSeek Reasoner', 'provider': 'deepseek'}
            ],
            'training_defaults': {
                'epochs': 10,
                'batch_size': 8,
                'learning_rate': 0.0001,
                'max_tokens': 512,
                'temperature': 0.7
            },
            'file_formats': {
                'training_guides': ['.pdf', '.doc', '.docx', '.txt', '.md', '.rtf', '.html', '.htm'],
                'test_submissions': ['.pdf', '.doc', '.docx', '.txt', '.jpg', '.jpeg', '.png', '.md', '.rtf', '.html', '.htm']
            },
            'limits': {
                'max_file_size_mb': 50,
                'max_files_per_upload': 10,
                'max_training_jobs': 5
            }
        }
        
        return jsonify({
            'success': True,
            'config': config
        })
        
    except Exception as e:
        logger.error(f"Error getting LLM training config: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to load configuration'
        }), 500

@llm_training_bp.route('/api/datasets')
@login_required
def get_datasets():
    """Get user's datasets"""
    try:
        datasets = LLMDataset.query.filter_by(user_id=current_user.id).all()
        return jsonify({
            'success': True,
            'datasets': [dataset.to_dict() for dataset in datasets]
        })
    except Exception as e:
        logger.error(f"Error fetching datasets: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/datasets', methods=['POST'])
@login_required
def create_dataset():
    """Create a new dataset"""
    try:
        data = request.get_json()
        
        # CSRF token is now handled by Flask-WTF automatically via X-CSRFToken header
        
        # Validate required fields
        if not data.get('name'):
            return jsonify({'success': False, 'error': 'Dataset name is required'}), 400
        
        # Create dataset
        dataset = LLMDataset(
            name=data['name'],
            description=data.get('description', ''),
            user_id=current_user.id,
            document_count=0,
            total_words=0,
            total_size=0
        )
        
        db.session.add(dataset)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'dataset': dataset.to_dict()
        })
        
    except Exception as e:
        logger.error(f"Error creating dataset: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/datasets/<dataset_id>', methods=['PUT'])
@login_required
def update_dataset(dataset_id):
    """Update a dataset"""
    try:
        # Find the dataset and verify ownership
        dataset = LLMDataset.query.filter_by(
            id=dataset_id,
            user_id=current_user.id
        ).first()
        
        if not dataset:
            return jsonify({'success': False, 'error': 'Dataset not found'}), 404
        
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'Request data required'}), 400
        
        # Update dataset properties
        if 'name' in data:
            name = data['name'].strip()
            if not name:
                return jsonify({'success': False, 'error': 'Dataset name cannot be empty'}), 400
            dataset.name = name
        
        if 'description' in data:
            dataset.description = data['description'].strip()
        
        db.session.commit()
        
        logger.info(f"Dataset {dataset_id} updated by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'dataset': dataset.to_dict(),
            'message': 'Dataset updated successfully'
        })
        
    except Exception as e:
        logger.error(f"Error updating dataset {dataset_id}: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/datasets/<dataset_id>', methods=['DELETE'])
@login_required
def delete_dataset(dataset_id):
    """Delete a dataset"""
    try:
        # Find the dataset and verify ownership
        dataset = LLMDataset.query.filter_by(
            id=dataset_id,
            user_id=current_user.id
        ).first()
        
        if not dataset:
            return jsonify({'success': False, 'error': 'Dataset not found'}), 404
        
        active_jobs = LLMTrainingJob.query.filter_by(
            dataset_id=dataset_id,
            user_id=current_user.id
        ).filter(
            LLMTrainingJob.status.in_(['pending', 'preparing', 'training', 'evaluating'])
        ).count()
        
        if active_jobs > 0:
            return jsonify({
                'success': False, 
                'error': f'Cannot delete dataset. It is being used by {active_jobs} active training job(s).'
            }), 400
        
        # Remove all document associations
        LLMDatasetDocument.query.filter_by(dataset_id=dataset_id).delete()
        
        # Delete the dataset
        db.session.delete(dataset)
        db.session.commit()
        
        logger.info(f"Dataset {dataset_id} ({dataset.name}) deleted by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': f'Dataset "{dataset.name}" deleted successfully'
        })
        
    except Exception as e:
        logger.error(f"Error deleting dataset {dataset_id}: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/datasets/cleanup', methods=['POST'])
@login_required
def cleanup_empty_datasets():
    """Clean up datasets that have no documents"""
    try:
        # Find datasets with no documents
        empty_datasets = LLMDataset.query.filter_by(
            user_id=current_user.id,
            document_count=0
        ).all()
        
        deleted_count = 0
        deleted_names = []
        
        for dataset in empty_datasets:
            # Double-check that there are really no documents
            doc_count = LLMDatasetDocument.query.filter_by(dataset_id=dataset.id).count()
            if doc_count == 0:
                deleted_names.append(dataset.name)
                db.session.delete(dataset)
                deleted_count += 1
        
        db.session.commit()
        
        logger.info(f"Cleaned up {deleted_count} empty datasets for user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': f'Cleaned up {deleted_count} empty dataset(s)',
            'deleted_datasets': deleted_names
        })
        
    except Exception as e:
        logger.error(f"Error cleaning up empty datasets: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/datasets/reset-counts', methods=['POST'])
@login_required
def reset_dataset_counts():
    """Reset dataset counts based on actual document associations"""
    try:
        datasets = LLMDataset.query.filter_by(user_id=current_user.id).all()
        updated_count = 0
        
        for dataset in datasets:
            # Get actual document associations
            dataset_docs = LLMDatasetDocument.query.filter_by(dataset_id=dataset.id).all()
            
            # Calculate actual counts
            actual_doc_count = len(dataset_docs)
            actual_total_size = 0
            actual_total_words = 0
            
            for dataset_doc in dataset_docs:
                document = LLMDocument.query.get(dataset_doc.document_id)
                if document:
                    actual_total_size += document.file_size or 0
                    actual_total_words += document.word_count or 0
            
            if (dataset.document_count != actual_doc_count or 
                dataset.total_size != actual_total_size or 
                dataset.total_words != actual_total_words):
                
                dataset.document_count = actual_doc_count
                dataset.total_size = actual_total_size
                dataset.total_words = actual_total_words
                updated_count += 1
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Updated counts for {updated_count} dataset(s)'
        })
        
    except Exception as e:
        logger.error(f"Error resetting dataset counts: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/training-jobs')
@login_required
def get_training_jobs():
    """Get user's training jobs"""
    try:
        jobs = LLMTrainingJob.query.filter_by(user_id=current_user.id).order_by(
            LLMTrainingJob.created_at.desc()
        ).all()
        
        return jsonify({
            'success': True,
            'jobs': [job.to_dict() for job in jobs]
        })
    except Exception as e:
        logger.error(f"Error fetching training jobs: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/training-jobs', methods=['POST'])
@login_required
def create_training_job():
    """Create a new training job with comprehensive validation"""
    try:
        from src.services.validation_service import ValidationService
        from src.services.error_handling_service import LLMTrainingErrorHandler
        
        data = request.get_json()
        
        # Enhanced validation
        if not data:
            return jsonify({'success': False, 'error': 'Request body is required'}), 400
        
        # Validate required fields with detailed messages
        required_fields = {
            'name': 'Training job name is required',
            'model': 'Model selection is required', 
            'dataset_id': 'Dataset selection is required'
        }
        
        for field, message in required_fields.items():
            if not data.get(field) or type(data.get(field)) is not str or not data.get(field).strip():
                return jsonify({'success': False, 'error': message}), 400
        
        # Prepare data for validation service - it expects model_id not model
        validation_data = data.copy()
        validation_data['model_id'] = validation_data.get('model')
        
        # Validate training configuration
        validation_service = ValidationService()
        config_validation = validation_service.validate_training_config(validation_data)
        
        if not config_validation['valid']:
            return jsonify({
                'success': False, 
                'error': 'Training configuration is invalid',
                'validation_errors': config_validation['errors'],
                'warnings': config_validation['warnings'],
                'recommendations': config_validation['recommendations']
            }), 400
        
        # Verify dataset exists and belongs to user
        dataset = LLMDataset.query.filter_by(
            id=data['dataset_id'], 
            user_id=current_user.id
        ).first()
        
        if not dataset:
            return jsonify({'success': False, 'error': 'Dataset not found or access denied'}), 404
        
        # Validate dataset integrity
        dataset_validation = validation_service.validate_dataset_integrity(data['dataset_id'])
        if not dataset_validation['valid']:
            return jsonify({
                'success': False,
                'error': 'Selected dataset is not suitable for training',
                'dataset_errors': dataset_validation['errors'],
                'dataset_warnings': dataset_validation['warnings'],
                'recommendations': dataset_validation['recommendations']
            }), 400
        
        normalized_config = config_validation['normalized_config']
        
        # Create training job with enhanced fields
        job = LLMTrainingJob(
            name=data['name'],
            model_id=data['model'],
            dataset_id=data['dataset_id'],
            user_id=current_user.id,
            status='pending',
            config_epochs=normalized_config.get('epochs', 10),
            config_batch_size=normalized_config.get('batch_size', 8),
            config_learning_rate=normalized_config.get('learning_rate', 0.0001),
            config_max_tokens=normalized_config.get('max_tokens', 512),
            config_temperature=data.get('temperature'),
            config_custom_parameters=data.get('custom_parameters', {}),
            validation_results={
                'config_validation': config_validation,
                'dataset_validation': dataset_validation,
                'validated_at': datetime.now(timezone.utc).isoformat()
            }
        )
        
        db.session.add(job)
        db.session.commit()
        
        # Start training asynchronously with enhanced error handling
        try:
            training_service = LLMTrainingService(current_app._get_current_object())
            training_service.start_training_async(job.id)
            
            response_data = {
                'success': True,
                'job': job.to_dict(),
                'message': 'Training job created and started successfully'
            }
            
            if config_validation.get('warnings') or dataset_validation.get('warnings'):
                response_data['warnings'] = (
                    config_validation.get('warnings', []) + 
                    dataset_validation.get('warnings', [])
                )
            
            return jsonify(response_data)
            
        except Exception as training_error:
            logger.warning(f"Could not start training automatically: {training_error}")
            
            # Update job status to indicate manual start required
            job.status = 'created'
            job.error_message = f"Automatic start failed: {str(training_error)}"
            db.session.commit()
            
            return jsonify({
                'success': True,
                'job': job.to_dict(),
                'message': 'Training job created successfully, but requires manual start',
                'warning': 'Automatic training start failed - use the Start button to begin training'
            })
        
    except Exception as e:
        logger.error(f"Error creating training job: {e}")
        db.session.rollback()
        
        # Use enhanced error handling
        error_response = LLMTrainingErrorHandler.handle_training_error(
            e, 'create_job', {'user_id': current_user.id, 'data': data}
        )
        
        return jsonify({
            'success': False,
            'error': error_response['error_message'],
            'error_type': error_response['error_type'],
            'recovery_suggestions': error_response.get('recovery_suggestions', []),
            'can_retry': True
        }), 500

@llm_training_bp.before_request
def log_request_info():
    """Log request information for debugging"""
    if request.endpoint and 'start_training_job' in request.endpoint:
        logger.info(f"Request to {request.endpoint}: {request.method} {request.url}")
        logger.info(f"Headers: {dict(request.headers)}")
        logger.info(f"Is JSON: {request.is_json}")
        logger.info(f"Content-Type: {request.content_type}")
        logger.info(f"User authenticated: {current_user.is_authenticated if hasattr(current_user, 'is_authenticated') else 'Unknown'}")

@llm_training_bp.route('/api/training-jobs/<job_id>/start', methods=['POST'])
@login_required
def start_training_job(job_id):
    """Start a training job with enhanced error handling and validation"""
    try:
        logger.info(f"Starting training job {job_id} for user {current_user.id}")
        
        # Get job with user verification
        job = LLMTrainingJob.query.filter_by(
            id=job_id, 
            user_id=current_user.id
        ).first()
        
        if not job:
            logger.warning(f"Training job {job_id} not found for user {current_user.id}")
            return jsonify({'success': False, 'error': 'Training job not found'}), 404
        
        logger.info(f"Job found: {job.name}, current status: {job.status}")
        
        # Validate job status and handle stuck jobs
        valid_start_statuses = ['pending', 'failed', 'cancelled', 'created']
        
        # Handle stuck jobs in various statuses
        if job.status in ['preparing', 'training', 'evaluating']:
            # Check if job has been stuck for more than 5 minutes
            time_threshold = 300  # 5 minutes in seconds
            current_time = datetime.now(timezone.utc)
            
            # Check against updated_at or created_at
            last_update_time = job.updated_at or job.created_at
            if last_update_time:
                if last_update_time.tzinfo is None:
                    last_update_time = last_update_time.replace(tzinfo=timezone.utc)
                
                time_since_update = current_time - last_update_time
                
                if time_since_update.total_seconds() > time_threshold:
                    logger.warning(f"Job {job_id} stuck in {job.status} status for {time_since_update.total_seconds()} seconds, resetting to failed")
                    job.status = 'failed'
                    job.error_message = f'Job was stuck in {job.status} status for {time_since_update.total_seconds():.0f} seconds and has been reset. You can now restart it.'
                    job.end_time = current_time
                    db.session.commit()
                else:
                    # Job is still within reasonable time, but user wants to restart
                    logger.info(f"Job {job_id} is in {job.status} status for {time_since_update.total_seconds()} seconds, allowing force restart")
                    job.status = 'failed'
                    job.error_message = f'Job was force-restarted from {job.status} status by user request'
                    job.end_time = current_time
                    db.session.commit()
            else:
                # No timestamp, assume it's stuck
                logger.warning(f"Job {job_id} in {job.status} status with no timestamp, resetting to failed")
                job.status = 'failed'
                job.error_message = f'Job was stuck in {job.status} status and has been reset'
                job.end_time = current_time
                db.session.commit()
        elif job.status not in valid_start_statuses:
            return jsonify({
                'success': False, 
                'error': f'Job cannot be started from status "{job.status}". Valid statuses: {", ".join(valid_start_statuses)}. If the job is stuck, try refreshing the page.'
            }), 400
        
        # Pre-start validation
        try:
            # Verify dataset exists and has content
            dataset = LLMDataset.query.filter_by(
                id=job.dataset_id, 
                user_id=current_user.id
            ).first()
            
            if not dataset:
                return jsonify({
                    'success': False, 
                    'error': 'Dataset not found or access denied'
                }), 400
            
            if dataset.document_count == 0:
                return jsonify({
                    'success': False, 
                    'error': 'Dataset has no documents. Add documents before starting training.'
                }), 400
            
            # Verify model is available
            from src.services.llm_training_service import LLMTrainingService
            training_service = LLMTrainingService(current_app._get_current_object())
            available_models = training_service.get_available_models()
            
            model_ids = [model['id'] for model in available_models]
            if job.model_id not in model_ids:
                return jsonify({
                    'success': False, 
                    'error': f'Model "{job.model_id}" is not available. Available models: {", ".join(model_ids)}'
                }), 400
            
        except Exception as validation_error:
            logger.warning(f"Pre-start validation failed for job {job_id}: {validation_error}")
            return jsonify({
                'success': False, 
                'error': f'Pre-start validation failed: {str(validation_error)}'
            }), 400
        
        # Update job status and clear any previous errors
        job.status = 'pending'
        job.error_message = None
        job.progress = 0
        job.current_epoch = 0
        
        # Set total epochs if not set
        if not job.total_epochs:
            job.total_epochs = job.config_epochs or 10
        
        db.session.commit()
        logger.info(f"Job {job_id} status updated to pending")
        
        try:
            # Start training asynchronously
            training_service.start_training_async(job_id)
            
            logger.info(f"Training job {job_id} started successfully")
            return jsonify({
                'success': True,
                'message': 'Training job started successfully',
                'job': {
                    'id': job.id,
                    'status': job.status,
                    'progress': job.progress,
                    'name': job.name
                }
            })
            
        except Exception as training_error:
            logger.error(f"Failed to start training for job {job_id}: {training_error}")
            logger.error(f"Training error traceback: {traceback.format_exc()}")
            
            # Update job status to indicate failure
            job.status = 'failed'
            job.error_message = f"Failed to start training: {str(training_error)}"
            db.session.commit()
            
            return jsonify({
                'success': False, 
                'error': f'Failed to start training: {str(training_error)}'
            }), 500
        
    except Exception as e:
        logger.error(f"Critical error in start_training_job route: {e}")
        logger.error(f"Exception type: {type(e).__name__}")
        logger.error(f"Exception traceback: {traceback.format_exc()}")
        
        # Try to update job status if we have the job object
        try:
            if 'job' in locals() and job:
                job.status = 'failed'
                job.error_message = f"Critical error: {str(e)}"
                db.session.commit()
        except:
            pass  # Don't let database update errors mask the original error
        
        return jsonify({
            'success': False, 
            'error': f'Internal server error: {str(e)}'
        }), 500

@llm_training_bp.route('/api/training-jobs/<job_id>/cancel', methods=['POST'])
@login_required
def cancel_training_job(job_id):
    """Cancel a training job"""
    try:
        job = LLMTrainingJob.query.filter_by(
            id=job_id, 
            user_id=current_user.id
        ).first()
        
        if not job:
            return jsonify({'success': False, 'error': 'Training job not found'}), 404
        
        if job.status not in ['pending', 'preparing', 'training']:
            return jsonify({'success': False, 'error': 'Job cannot be cancelled in current state'}), 400
        
        # Cancel training
        training_service = LLMTrainingService(current_app)
        training_service.cancel_training(job_id)
        
        return jsonify({'success': True})
        
    except Exception as e:
        logger.error(f"Error cancelling training job: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/training-jobs/<job_id>', methods=['DELETE'])
@login_required
def delete_training_job(job_id):
    """Delete a training job"""
    try:
        job = LLMTrainingJob.query.filter_by(
            id=job_id, 
            user_id=current_user.id
        ).first()
        
        if not job:
            return jsonify({'success': False, 'error': 'Training job not found'}), 404
        
        # Check if job is actually running or just stuck
        if job.status in ['training', 'preparing', 'evaluating']:
            # Check if job has been stuck for more than 5 minutes
            time_threshold = 300  # 5 minutes in seconds
            current_time = datetime.now(timezone.utc)
            
            last_update_time = job.updated_at or job.created_at
            if last_update_time:
                if last_update_time.tzinfo is None:
                    last_update_time = last_update_time.replace(tzinfo=timezone.utc)
                
                time_since_update = current_time - last_update_time
                
                if time_since_update.total_seconds() > time_threshold:
                    # Job is stuck, allow deletion
                    logger.warning(f"Allowing deletion of stuck job {job_id} in {job.status} status for {time_since_update.total_seconds()} seconds")
                else:
                    # Job might still be running
                    return jsonify({
                        'success': False, 
                        'error': f'Cannot delete a job in {job.status} status. Please wait for it to complete or try again in a few minutes if it appears stuck.'
                    }), 400
            # If no timestamp, assume it's stuck and allow deletion
        
        job_name = job.name
        
        # Delete the training job
        db.session.delete(job)
        db.session.commit()
        
        logger.info(f"Training job {job_id} ({job_name}) deleted by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': f'Training job "{job_name}" deleted successfully'
        })
        
    except Exception as e:
        logger.error(f"Error deleting training job {job_id}: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/training-jobs/<job_id>/status')
@login_required
def get_training_job_status(job_id):
    """Get detailed status of a specific training job"""
    try:
        job = LLMTrainingJob.query.filter_by(
            id=job_id, 
            user_id=current_user.id
        ).first()
        
        if not job:
            return jsonify({'success': False, 'error': 'Training job not found'}), 404
        
        # Get detailed job information
        job_data = {
            'id': job.id,
            'name': job.name,
            'status': job.status,
            'progress': job.progress or 0,
            'current_epoch': job.current_epoch or 0,
            'total_epochs': job.total_epochs or 0,
            'accuracy': job.accuracy or 0.0,
            'loss': job.loss or 0.0,
            'validation_accuracy': job.validation_accuracy or 0.0,
            'model_id': job.model_id,
            'dataset_id': job.dataset_id,
            'created_at': job.created_at.isoformat() if job.created_at else None,
            'start_time': job.start_time.isoformat() if job.start_time else None,
            'end_time': job.end_time.isoformat() if job.end_time else None,
            'error_message': job.error_message,
            'evaluation_results': job.evaluation_results or {},
            'training_metrics': job.training_metrics or {},
            'health_metrics': job.health_metrics or {}
        }
        
        # Add runtime calculation if training is active
        if job.start_time and job.status in ['preparing', 'training', 'evaluating']:
            from datetime import datetime, timezone
            current_time = datetime.now(timezone.utc)
            start_time = job.start_time
            if start_time.tzinfo is None:
                start_time = start_time.replace(tzinfo=timezone.utc)
            runtime_seconds = (current_time - start_time).total_seconds()
            job_data['runtime_minutes'] = runtime_seconds / 60
        
        return jsonify({
            'success': True,
            'job': job_data
        })
        
    except Exception as e:
        logger.error(f"Error getting training job status {job_id}: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/reports')
@login_required
def get_reports():
    """Get user's training reports"""
    try:
        reports = LLMTrainingReport.query.filter_by(user_id=current_user.id).order_by(
            LLMTrainingReport.created_at.desc()
        ).all()
        
        return jsonify({
            'success': True,
            'reports': [report.to_dict() for report in reports]
        })
    except Exception as e:
        logger.error(f"Error fetching reports: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/reports', methods=['POST'])
@login_required
def generate_report():
    """Generate a training report"""
    try:
        data = request.get_json()
        
        # Validate job IDs
        job_ids = data.get('job_ids', [])
        if not job_ids:
            return jsonify({'success': False, 'error': 'At least one job must be selected'}), 400
        
        # Verify jobs exist and belong to user
        jobs = LLMTrainingJob.query.filter(
            LLMTrainingJob.id.in_(job_ids),
            LLMTrainingJob.user_id == current_user.id
        ).all()
        
        if len(jobs) != len(job_ids):
            return jsonify({'success': False, 'error': 'Some jobs not found'}), 404
        
        # Create report
        report = LLMTrainingReport(
            name=f"Training Report - {len(jobs)} jobs",
            user_id=current_user.id,
            job_ids=job_ids,
            format=data.get('format', 'html'),
            include_metrics=data.get('include_metrics', True),
            include_logs=data.get('include_logs', False),
            status='generating'
        )
        
        db.session.add(report)
        db.session.commit()
        
        try:
            training_service = LLMTrainingService(current_app)
            training_service.generate_report_async(report.id)
        except Exception as report_error:
            logger.warning(f"Could not generate report automatically: {report_error}")
            # Report is still created, just not generated automatically
        
        return jsonify({
            'success': True,
            'report': report.to_dict()
        })
        
    except Exception as e:
        logger.error(f"Error generating report: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/reports/<report_id>', methods=['DELETE'])
@login_required
def delete_report(report_id):
    """Delete a training report"""
    try:
        # Find the report and verify ownership
        report = LLMTrainingReport.query.filter_by(
            id=report_id,
            user_id=current_user.id
        ).first()
        
        if not report:
            return jsonify({'success': False, 'error': 'Report not found'}), 404
        
        report_name = report.name
        
        if report.file_path and os.path.exists(report.file_path):
            try:
                os.remove(report.file_path)
            except OSError as e:
                logger.warning(f"Could not delete report file {report.file_path}: {e}")
        
        # Delete the report record
        db.session.delete(report)
        db.session.commit()
        
        logger.info(f"Report {report_id} ({report_name}) deleted by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': f'Report "{report_name}" deleted successfully'
        })
        
    except Exception as e:
        logger.error(f"Error deleting report {report_id}: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/reports/<report_id>/download')
@login_required
def download_report(report_id):
    """Download a training report"""
    try:
        # Find the report and verify ownership
        report = LLMTrainingReport.query.filter_by(
            id=report_id,
            user_id=current_user.id
        ).first()
        
        if not report:
            return jsonify({'success': False, 'error': 'Report not found'}), 404
        
        if report.status != 'completed':
            return jsonify({'success': False, 'error': 'Report is not ready for download'}), 400
        
        if not report.file_path or not os.path.exists(report.file_path):
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Training Report - {report.name}</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; }}
                    .header {{ background: #f8f9fa; padding: 20px; border-radius: 8px; margin-bottom: 20px; }}
                    .section {{ margin-bottom: 30px; }}
                    .job {{ background: #f8f9fa; padding: 15px; margin: 10px 0; border-radius: 5px; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>Training Report: {report.name}</h1>
                    <p>Generated: {report.created_at.strftime('%Y-%m-%d %H:%M:%S')}</p>
                    <p>Format: {report.format.upper()}</p>
                </div>
                
                <div class="section">
                    <h2>Report Summary</h2>
                    <p>This report covers {len(report.job_ids)} training job(s).</p>
                    <p>Job IDs: {', '.join(report.job_ids)}</p>
                </div>
                
                <div class="section">
                    <h2>Configuration</h2>
                    <ul>
                        <li>Include Metrics: {'Yes' if report.include_metrics else 'No'}</li>
                        <li>Include Logs: {'Yes' if report.include_logs else 'No'}</li>
                    </ul>
                </div>
                
                <div class="section">
                    <p><em>Note: This is a placeholder report. Full report generation is not yet implemented.</em></p>
                </div>
            </body>
            </html>
            """
            
            from flask import Response
            return Response(
                html_content,
                mimetype='text/html',
                headers={
                    'Content-Disposition': f'attachment; filename="training_report_{report_id}.html"'
                }
            )
        
        # Serve the actual file
        from flask import send_file
        return send_file(
            report.file_path,
            as_attachment=True,
            download_name=f"training_report_{report_id}.{report.format}"
        )
        
    except Exception as e:
        logger.error(f"Error downloading report {report_id}: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/documents')
@login_required
def get_documents():
    """Get user's documents"""
    try:
        documents = LLMDocument.query.filter_by(user_id=current_user.id).all()
        return jsonify({
            'success': True,
            'documents': [doc.to_dict() for doc in documents]
        })
    except Exception as e:
        logger.error(f"Error fetching documents: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/documents/upload', methods=['POST'])
@login_required
def upload_documents():
    """Upload documents for training with duplicate prevention"""
    try:
        if 'files' not in request.files:
            return jsonify({'success': False, 'error': 'No files provided'}), 400
        
        files = request.files.getlist('files')
        if not files or all(f.filename == '' for f in files):
            return jsonify({'success': False, 'error': 'No files selected'}), 400
        
        uploaded_documents = []
        skipped_files = []
        errors = []
        
        for file in files:
            if file and file.filename:
                try:
                    # Generate unique filename
                    import uuid
                    import os
                    from werkzeug.utils import secure_filename
                    import hashlib
                    
                    original_name = secure_filename(file.filename)
                    if not original_name:
                        errors.append(f"Invalid filename: {file.filename}")
                        continue
                    
                    file_extension = os.path.splitext(original_name)[1].lower()
                    
                    # Check for existing document with same name and user
                    existing_doc = LLMDocument.query.filter_by(
                        user_id=current_user.id,
                        original_name=original_name
                    ).first()
                    
                    if existing_doc:
                        skipped_files.append({
                            'name': original_name,
                            'reason': 'File already exists'
                        })
                        continue
                    
                    # Read file content for hash calculation
                    file.seek(0)  # Reset file pointer
                    file_content = file.read()
                    file.seek(0)  # Reset again for saving
                    
                    # Calculate file hash to detect duplicates
                    file_hash = hashlib.md5(file_content).hexdigest()
                    
                    # Check for files with similar size (basic duplicate detection)
                    file_size_bytes = len(file_content)
                    existing_size_docs = LLMDocument.query.filter_by(
                        user_id=current_user.id,
                        file_size=file_size_bytes
                    ).all()
                    
                    # If we find files with same size, do a more detailed check
                    is_duplicate = False
                    for existing_doc in existing_size_docs:
                        if existing_doc.original_name != original_name:  # Different name but same size
                            # This is a basic duplicate detection - could be enhanced with actual content comparison
                            is_duplicate = True
                            skipped_files.append({
                                'name': original_name,
                                'reason': f'Similar file already exists ({existing_doc.original_name})'
                            })
                            break
                    
                    if is_duplicate:
                        continue
                    
                    stored_name = f"{uuid.uuid4()}{file_extension}"
                    
                    upload_dir = os.path.join(current_app.root_path, 'uploads', 'llm_documents')
                    os.makedirs(upload_dir, exist_ok=True)
                    
                    file_path = os.path.join(upload_dir, stored_name)
                    
                    # Save file
                    with open(file_path, 'wb') as f:
                        f.write(file_content)
                    
                    # Get file info
                    file_size = os.path.getsize(file_path)
                    
                    text_content = ""
                    word_count = 0
                    
                    try:
                        text_content = extract_text_from_file(file_path, file_extension)
                        if text_content:
                            word_count = len(text_content.split())
                            logger.info(f"Extracted {word_count} words from {original_name}")
                        else:
                            logger.warning(f"No text content extracted from {original_name}")
                    except Exception as e:
                        logger.error(f"Error extracting text from {original_name}: {e}")
                        text_content = ""
                        word_count = 0
                    
                    # Create document record
                    document = LLMDocument(
                        user_id=current_user.id,
                        name=original_name,
                        original_name=original_name,
                        stored_name=stored_name,
                        file_type=file_extension[1:] if file_extension else 'unknown',
                        mime_type=file.mimetype or 'application/octet-stream',
                        file_size=file_size,
                        file_path=file_path,
                        text_content=text_content,
                        word_count=word_count,
                        character_count=len(text_content),
                        extracted_text=bool(text_content)
                    )
                    
                    db.session.add(document)
                    uploaded_documents.append(document)
                    
                except Exception as e:
                    logger.error(f"Error processing file {file.filename}: {e}")
                    errors.append(f"Error processing {file.filename}: {str(e)}")
                    continue
        
        # Commit all successful uploads
        if uploaded_documents:
            db.session.commit()
            logger.info(f"Successfully uploaded {len(uploaded_documents)} documents for user {current_user.id}")
        
        response_data = {
            'success': True,
            'documents': [doc.to_dict() for doc in uploaded_documents],
            'message': f'Successfully uploaded {len(uploaded_documents)} document(s)'
        }
        
        # Add warnings for skipped files
        if skipped_files:
            response_data['skipped_files'] = skipped_files
            response_data['message'] += f', skipped {len(skipped_files)} duplicate(s)'
        
        # Add errors if any
        if errors:
            response_data['errors'] = errors
            if not uploaded_documents:
                response_data['success'] = False
                response_data['message'] = 'No documents were uploaded due to errors'
        
        return jsonify(response_data)
        
    except Exception as e:
        logger.error(f"Error uploading documents: {e}")
        logger.error(traceback.format_exc())
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

# Model Testing API Endpoints

@llm_training_bp.route('/api/model-tests', methods=['POST'])
@login_required
def create_model_test():
    """Create a new model test session"""
    try:
        from src.services.model_testing_service import ModelTestingService
        
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['training_job_id', 'name']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'success': False, 'error': f'{field} is required'}), 400
        
        # Create test session
        testing_service = ModelTestingService()
        test_id = testing_service.create_test_session(
            user_id=current_user.id,
            training_job_id=data['training_job_id'],
            config=data
        )
        
        return jsonify({
            'success': True,
            'test_id': test_id,
            'message': 'Model test session created successfully'
        })
        
    except Exception as e:
        logger.error(f"Error creating model test: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/model-tests/<test_id>/submissions', methods=['POST'])
@login_required
def upload_test_submissions(test_id):
    """Upload test submissions for a model test"""
    try:
        from src.services.model_testing_service import ModelTestingService
        
        if 'files' not in request.files:
            return jsonify({'success': False, 'error': 'No files provided'}), 400
        
        files = request.files.getlist('files')
        if not files or all(f.filename == '' for f in files):
            return jsonify({'success': False, 'error': 'No files selected'}), 400
        
        # Process files
        file_data = []
        for file in files:
            if file and file.filename:
                file_info = {
                    'original_name': file.filename,
                    'content': file.read(),
                    'expected_grade': request.form.get(f'expected_grade_{file.filename}'),
                    'expected_feedback': request.form.get(f'expected_feedback_{file.filename}', '')
                }
                file_data.append(file_info)
        
        # Upload submissions
        testing_service = ModelTestingService()
        submissions = testing_service.upload_test_submissions(test_id, file_data)
        
        return jsonify({
            'success': True,
            'submissions': submissions,
            'message': f'Successfully uploaded {len(submissions)} test submission(s)'
        })
        
    except Exception as e:
        logger.error(f"Error uploading test submissions: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/model-tests/<test_id>/run', methods=['POST'])
@login_required
def run_model_test(test_id):
    """Start running a model test"""
    try:
        from src.services.model_testing_service import ModelTestingService
        
        testing_service = ModelTestingService()
        result = testing_service.run_model_test(test_id)
        
        return jsonify({
            'success': True,
            'result': result,
            'message': 'Model test started successfully'
        })
        
    except Exception as e:
        logger.error(f"Error running model test: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/model-tests/<test_id>/cancel', methods=['POST'])
@login_required
def cancel_model_test(test_id):
    """Cancel a running model test"""
    try:
        from src.services.model_testing_service import ModelTestingService
        
        testing_service = ModelTestingService()
        success = testing_service.cancel_test(test_id, current_user.id)
        
        if success:
            return jsonify({
                'success': True,
                'message': 'Model test cancelled successfully'
            })
        else:
            return jsonify({
                'success': False,
                'error': 'Failed to cancel model test'
            }), 400
        
    except Exception as e:
        logger.error(f"Error cancelling model test: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/model-tests/<test_id>/results')
@login_required
def get_model_test_results(test_id):
    """Get results for a completed model test"""
    try:
        from src.services.model_testing_service import ModelTestingService
        
        testing_service = ModelTestingService()
        results = testing_service.get_test_results(test_id, current_user.id)
        
        return jsonify({
            'success': True,
            'results': results
        })
        
    except Exception as e:
        logger.error(f"Error getting model test results: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/model-tests')
@login_required
def get_model_tests():
    """Get user's model tests"""
    try:
        from src.database.models import LLMModelTest
        
        tests = LLMModelTest.query.filter_by(user_id=current_user.id).order_by(
            LLMModelTest.created_at.desc()
        ).all()
        
        return jsonify({
            'success': True,
            'tests': [test.to_dict() for test in tests]
        })
        
    except Exception as e:
        logger.error(f"Error fetching model tests: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/model-tests/<test_id>', methods=['DELETE'])
@login_required
def delete_model_test(test_id):
    """Delete a model test"""
    try:
        from src.database.models import LLMModelTest
        
        # Find the test and verify ownership
        test = LLMModelTest.query.filter_by(
            id=test_id,
            user_id=current_user.id
        ).first()
        
        if not test:
            return jsonify({'success': False, 'error': 'Model test not found'}), 404
        
        if test.status in ['running']:
            return jsonify({
                'success': False, 
                'error': 'Cannot delete a running test. Please cancel it first.'
            }), 400
        
        test_name = test.name
        
        # Delete the test (cascade will handle submissions)
        db.session.delete(test)
        db.session.commit()
        
        logger.info(f"Model test {test_id} ({test_name}) deleted by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': f'Model test "{test_name}" deleted successfully'
        })
        
    except Exception as e:
        logger.error(f"Error deleting model test {test_id}: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/model-tests/<test_id>/report')
@login_required
def generate_model_test_report(test_id):
    """Generate a comprehensive report for a model test"""
    try:
        from src.services.model_testing_service import ModelTestingService
        
        testing_service = ModelTestingService()
        report = testing_service.generate_test_report(test_id, current_user.id)
        
        return jsonify({
            'success': True,
            'report': report
        })
        
    except Exception as e:
        logger.error(f"Error generating model test report: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/model-tests/<test_id>/status')
@login_required
def get_model_test_status(test_id):
    """Get current status of a model test"""
    try:
        from src.database.models import LLMModelTest
        
        test = LLMModelTest.query.filter_by(
            id=test_id,
            user_id=current_user.id
        ).first()
        
        if not test:
            return jsonify({'success': False, 'error': 'Model test not found'}), 404
        
        return jsonify({
            'success': True,
            'status': {
                'id': test.id,
                'status': test.status,
                'progress': test.progress,
                'processed_submissions': test.processed_submissions,
                'total_submissions': test.total_submissions,
                'started_at': test.started_at.isoformat() if test.started_at else None,
                'completed_at': test.completed_at.isoformat() if test.completed_at else None,
                'error_message': test.error_message
            }
        })
        
    except Exception as e:
        logger.error(f"Error getting model test status: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/documents/<document_id>', methods=['PUT'])
@login_required
def update_document(document_id):
    """Update a document"""
    try:
        # Find the document and verify ownership
        document = LLMDocument.query.filter_by(
            id=document_id, 
            user_id=current_user.id
        ).first()
        
        if not document:
            return jsonify({'success': False, 'error': 'Document not found'}), 404
        
        data = request.get_json()
        
        if 'name' in data:
            document.name = data['name']
        
        # Handle dataset assignment
        if 'dataset_id' in data:
            LLMDatasetDocument.query.filter_by(document_id=document_id).delete()
            
            if data['dataset_id']:
                # Verify dataset exists and belongs to user
                dataset = LLMDataset.query.filter_by(
                    id=data['dataset_id'],
                    user_id=current_user.id
                ).first()
                
                if not dataset:
                    return jsonify({'success': False, 'error': 'Dataset not found'}), 404
                
                # Create new dataset-document association
                dataset_doc = LLMDatasetDocument(
                    dataset_id=data['dataset_id'],
                    document_id=document_id
                )
                db.session.add(dataset_doc)
                
                # Update dataset counts
                dataset.document_count = LLMDatasetDocument.query.filter_by(dataset_id=data['dataset_id']).count() + 1
                dataset.total_size += document.file_size
                dataset.total_words += document.word_count
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'message': 'Document updated successfully',
            'document': document.to_dict()
        })
        
    except Exception as e:
        logger.error(f"Error updating document: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/datasets/<dataset_id>/documents', methods=['POST'])
@login_required
def add_documents_to_dataset(dataset_id):
    """Add documents to a dataset"""
    try:
        # Verify dataset exists and belongs to user
        dataset = LLMDataset.query.filter_by(
            id=dataset_id,
            user_id=current_user.id
        ).first()
        
        if not dataset:
            return jsonify({'success': False, 'error': 'Dataset not found'}), 404
        
        data = request.get_json()
        if not data or 'document_ids' not in data:
            return jsonify({'success': False, 'error': 'Document IDs are required'}), 400
        
        document_ids = data['document_ids']
        if not isinstance(document_ids, list) or not document_ids:
            return jsonify({'success': False, 'error': 'At least one document ID is required'}), 400
        
        # Verify all documents exist and belong to user
        documents = LLMDocument.query.filter(
            LLMDocument.id.in_(document_ids),
            LLMDocument.user_id == current_user.id
        ).all()
        
        if len(documents) != len(document_ids):
            return jsonify({'success': False, 'error': 'Some documents not found or access denied'}), 404
        
        # Check for existing associations
        existing_associations = LLMDatasetDocument.query.filter(
            LLMDatasetDocument.dataset_id == dataset_id,
            LLMDatasetDocument.document_id.in_(document_ids)
        ).all()
        
        existing_doc_ids = {assoc.document_id for assoc in existing_associations}
        new_document_ids = [doc_id for doc_id in document_ids if doc_id not in existing_doc_ids]
        
        # Add new associations
        added_count = 0
        total_size_added = 0
        total_words_added = 0
        
        for document in documents:
            if document.id in new_document_ids:
                dataset_doc = LLMDatasetDocument(
                    dataset_id=dataset_id,
                    document_id=document.id
                )
                db.session.add(dataset_doc)
                added_count += 1
                total_size_added += document.file_size or 0
                total_words_added += document.word_count or 0
        
        # Update dataset counts
        dataset.document_count += added_count
        dataset.total_size += total_size_added
        dataset.total_words += total_words_added
        
        db.session.commit()
        
        logger.info(f"Added {added_count} documents to dataset {dataset_id} by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': f'Successfully added {added_count} document(s) to dataset',
            'added_count': added_count,
            'skipped_count': len(document_ids) - added_count,
            'dataset': dataset.to_dict()
        })
        
    except Exception as e:
        logger.error(f"Error adding documents to dataset {dataset_id}: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/datasets/<dataset_id>/documents/<document_id>', methods=['DELETE'])
@login_required
def remove_document_from_dataset(dataset_id, document_id):
    """Remove a document from a dataset"""
    try:
        # Verify dataset exists and belongs to user
        dataset = LLMDataset.query.filter_by(
            id=dataset_id,
            user_id=current_user.id
        ).first()
        
        if not dataset:
            return jsonify({'success': False, 'error': 'Dataset not found'}), 404
        
        # Verify document exists and belongs to user
        document = LLMDocument.query.filter_by(
            id=document_id,
            user_id=current_user.id
        ).first()
        
        if not document:
            return jsonify({'success': False, 'error': 'Document not found'}), 404
        
        # Find and remove the association
        association = LLMDatasetDocument.query.filter_by(
            dataset_id=dataset_id,
            document_id=document_id
        ).first()
        
        if not association:
            return jsonify({'success': False, 'error': 'Document is not in this dataset'}), 404
        
        # Remove association
        db.session.delete(association)
        
        # Update dataset counts
        dataset.document_count = max(0, dataset.document_count - 1)
        dataset.total_size = max(0, dataset.total_size - (document.file_size or 0))
        dataset.total_words = max(0, dataset.total_words - (document.word_count or 0))
        
        db.session.commit()
        
        logger.info(f"Removed document {document_id} from dataset {dataset_id} by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': f'Document "{document.name}" removed from dataset',
            'dataset': dataset.to_dict()
        })
        
    except Exception as e:
        logger.error(f"Error removing document from dataset: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/datasets/<dataset_id>/documents')
@login_required
def get_dataset_documents(dataset_id):
    """Get documents in a dataset"""
    try:
        # Verify dataset exists and belongs to user
        dataset = LLMDataset.query.filter_by(
            id=dataset_id,
            user_id=current_user.id
        ).first()
        
        if not dataset:
            return jsonify({'success': False, 'error': 'Dataset not found'}), 404
        
        # Get documents in the dataset
        documents = db.session.query(LLMDocument).join(LLMDatasetDocument).filter(
            LLMDatasetDocument.dataset_id == dataset_id,
            LLMDocument.user_id == current_user.id
        ).all()
        
        return jsonify({
            'success': True,
            'documents': [doc.to_dict() for doc in documents],
            'dataset': dataset.to_dict()
        })
        
    except Exception as e:
        logger.error(f"Error getting dataset documents: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/documents/<document_id>', methods=['DELETE'])
@login_required
def delete_document(document_id):
    """Delete a document"""
    try:
        # Find the document and verify ownership
        document = LLMDocument.query.filter_by(
            id=document_id, 
            user_id=current_user.id
        ).first()
        
        if not document:
            return jsonify({'success': False, 'error': 'Document not found'}), 404
        
        LLMDatasetDocument.query.filter_by(document_id=document_id).delete()
        
        # Update dataset counts
        datasets_to_update = db.session.query(LLMDataset).join(LLMDatasetDocument).filter(
            LLMDatasetDocument.document_id == document_id
        ).all()
        
        for dataset in datasets_to_update:
            dataset.document_count = max(0, dataset.document_count - 1)
            dataset.total_size = max(0, dataset.total_size - document.file_size)
            dataset.total_words = max(0, dataset.total_words - document.word_count)
        
        import os
        if document.file_path and os.path.exists(document.file_path):
            try:
                os.remove(document.file_path)
            except OSError as e:
                logger.warning(f"Could not delete file {document.file_path}: {e}")
        
        # Delete the document record
        db.session.delete(document)
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Document deleted successfully'})
        
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/models/refresh', methods=['POST'])
@login_required
def refresh_models():
    """Refresh available models list"""
    try:
        training_service = LLMTrainingService(current_app)
        models = training_service.get_available_models()
        
        return jsonify({
            'success': True,
            'models': models,
            'message': f'Found {len(models)} available models'
        })
        
    except Exception as e:
        logger.error(f"Error refreshing models: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/models')
@login_required
def get_models():
    """Get available models list"""
    try:
        training_service = LLMTrainingService(current_app)
        models = training_service.get_available_models()
        
        return jsonify({
            'success': True,
            'models': models
        })
        
    except Exception as e:
        logger.error(f"Error getting models: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/test', methods=['POST'])
@login_required
def test_functionality():
    """Test LLM training functionality"""
    try:
        training_service = LLMTrainingService(current_app)
        test_results = training_service.run_system_test()
        
        return jsonify({
            'success': True,
            'test_results': test_results,
            'message': 'System test completed successfully'
        })
        
    except Exception as e:
        logger.error(f"Error running system test: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@llm_training_bp.route('/api/stats')
@login_required
def get_stats():
    """Get training statistics for the user"""
    try:
        # Dataset stats
        total_datasets = LLMDataset.query.filter_by(user_id=current_user.id).count()
        
        total_documents = LLMDocument.query.filter_by(user_id=current_user.id).count()
        total_size = db.session.query(db.func.sum(LLMDocument.file_size)).filter_by(
            user_id=current_user.id
        ).scalar() or 0
        total_words = db.session.query(db.func.sum(LLMDocument.word_count)).filter_by(
            user_id=current_user.id
        ).scalar() or 0
        
        # Training job stats
        total_jobs = LLMTrainingJob.query.filter_by(user_id=current_user.id).count()
        running_jobs = LLMTrainingJob.query.filter_by(
            user_id=current_user.id, 
            status='training'
        ).count()
        completed_jobs = LLMTrainingJob.query.filter_by(
            user_id=current_user.id, 
            status='completed'
        ).count()
        failed_jobs = LLMTrainingJob.query.filter_by(
            user_id=current_user.id, 
            status='failed'
        ).count()
        
        return jsonify({
            'success': True,
            'stats': {
                'datasets': {
                    'total_datasets': total_datasets,
                    'total_documents': total_documents,
                    'total_words': total_words,
                    'total_size_mb': round(total_size / (1024 * 1024), 2) if total_size else 0
                },
                'training': {
                    'total_jobs': total_jobs,
                    'running_jobs': running_jobs,
                    'completed_jobs': completed_jobs,
                    'failed_jobs': failed_jobs
                }
            }
        })
        
    except Exception as e:
        logger.error(f"Error fetching stats: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500
# Enhanced LLM Training Workflow Routes

@llm_training_bp.route('/api/training-guides', methods=['GET'])
@login_required
def get_training_guides():
    """Get all training guides for the current user"""
    try:
        # For now, we'll use the existing LLMDocument model to store training guides
        # with a special type field to distinguish them
        guides = LLMDocument.query.filter_by(
            user_id=current_user.id,
            type='training_guide'
        ).order_by(LLMDocument.created_at.desc()).all()
        
        guides_data = []
        for guide in guides:
            guides_data.append({
                'id': guide.id,
                'name': guide.name,
                'description': '',  # Description not stored in model anymore
                'file_size': guide.file_size,
                'word_count': guide.word_count or 0,
                'created_at': guide.created_at.isoformat(),
                'file_path': guide.file_path
            })
        
        return jsonify({
            'success': True,
            'guides': guides_data
        })
        
    except Exception as e:
        logger.error(f"Error getting training guides: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to load training guides'
        }), 500

@llm_training_bp.route('/api/training-guides/upload', methods=['POST'])
@login_required
def upload_training_guide():
    """Upload a new training guide"""
    try:
        name = request.form.get('name', '').strip()
        description = request.form.get('description', '').strip()
        
        if not name:
            return jsonify({
                'success': False,
                'error': 'Guide name is required'
            }), 400
        
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'error': 'No file provided'
            }), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': 'No file selected'
            }), 400
        
        # Validate file type
        allowed_extensions = {'.pdf', '.doc', '.docx', '.txt', '.md', '.rtf', '.html', '.htm'}
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension not in allowed_extensions:
            return jsonify({
                'success': False,
                'error': f'File type {file_extension} not supported. Allowed types: {", ".join(allowed_extensions)}'
            }), 400
        
        # Create upload directory
        upload_dir = os.path.join(current_app.root_path, 'uploads', 'training_guides')
        os.makedirs(upload_dir, exist_ok=True)
        
        # Generate unique filename
        filename = f"{current_user.id}_{int(datetime.now().timestamp())}_{file.filename}"
        file_path = os.path.join(upload_dir, filename)
        
        # Save file
        file.save(file_path)
        
        # Extract text content
        try:
            text_content = extract_text_from_file(file_path, file_extension)
            word_count = len(text_content.split()) if text_content else 0
        except Exception as e:
            logger.warning(f"Failed to extract text from training guide {filename}: {e}")
            text_content = ""
            word_count = 0
        
        # Check for duplicate content
        from src.utils.content_deduplication import check_llm_document_duplicate, get_deduplication_response, calculate_content_hash
        
        is_duplicate, existing_doc = check_llm_document_duplicate(
            user_id=current_user.id,
            content=text_content,
            document_type='training_guide',
            db_session=db.session
        )
        
        if is_duplicate:
            # Remove the uploaded file since it's a duplicate
            try:
                os.remove(file_path)
            except OSError:
                pass
            
            return jsonify(get_deduplication_response(existing_doc, "training guide")), 409
        
        # Create database record
        guide = LLMDocument(
            user_id=current_user.id,
            name=name,
            original_name=file.filename,
            stored_name=filename,
            file_type=file_extension,
            mime_type=file.mimetype or 'application/octet-stream',
            file_size=os.path.getsize(file_path),
            file_path=file_path,
            text_content=text_content,
            content_hash=calculate_content_hash(text_content),
            word_count=word_count,
            character_count=len(text_content) if text_content else 0,
            extracted_text=bool(text_content),
            type='training_guide'
        )
        
        db.session.add(guide)
        db.session.commit()
        
        logger.info(f"Training guide uploaded successfully: {name} by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': 'Training guide uploaded successfully',
            'guide': {
                'id': guide.id,
                'name': guide.name,
                'description': description,
                'word_count': word_count
            }
        })
        
    except Exception as e:
        logger.error(f"Error uploading training guide: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to upload training guide'
        }), 500

@llm_training_bp.route('/api/training-guides/<guide_id>', methods=['DELETE'])
@login_required
def delete_training_guide(guide_id):
    """Delete a training guide"""
    try:
        guide = LLMDocument.query.filter_by(
            id=guide_id,
            user_id=current_user.id,
            type='training_guide'
        ).first()
        
        if not guide:
            return jsonify({
                'success': False,
                'error': 'Training guide not found'
            }), 404
        
        # Delete file
        try:
            if os.path.exists(guide.file_path):
                os.remove(guide.file_path)
        except Exception as e:
            logger.warning(f"Failed to delete training guide file {guide.file_path}: {e}")
        
        # Delete database record
        db.session.delete(guide)
        db.session.commit()
        
        logger.info(f"Training guide deleted: {guide.name} by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': 'Training guide deleted successfully'
        })
        
    except Exception as e:
        logger.error(f"Error deleting training guide: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to delete training guide'
        }), 500

@llm_training_bp.route('/api/test-submissions', methods=['GET'])
@login_required
def get_test_submissions():
    """Get all test submissions for the current user"""
    try:
        submissions = LLMDocument.query.filter_by(
            user_id=current_user.id,
            type='test_submission'
        ).order_by(LLMDocument.created_at.desc()).all()
        
        submissions_data = []
        for submission in submissions:
            submissions_data.append({
                'id': submission.id,
                'name': submission.name,
                'description': '',  # Description not stored in model anymore
                'size': submission.file_size,  # Use file_size instead of size
                'expected_score': None,  # Expected score not stored in model anymore
                'created_at': submission.created_at.isoformat(),
                'file_path': submission.file_path
            })
        
        return jsonify({
            'success': True,
            'submissions': submissions_data
        })
        
    except Exception as e:
        logger.error(f"Error getting test submissions: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to load test submissions'
        }), 500

@llm_training_bp.route('/api/test-submissions/upload', methods=['POST'])
@login_required
def upload_test_submission():
    """Upload a new test submission"""
    try:
        name = request.form.get('name', '').strip()
        description = request.form.get('description', '').strip()
        expected_score = request.form.get('expected_score')
        
        if not name:
            return jsonify({
                'success': False,
                'error': 'Submission name is required'
            }), 400
        
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'error': 'No file provided'
            }), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': 'No file selected'
            }), 400
        
        # Validate expected score
        if expected_score:
            try:
                expected_score = float(expected_score)
                if expected_score < 0 or expected_score > 100:
                    return jsonify({
                        'success': False,
                        'error': 'Expected score must be between 0 and 100'
                    }), 400
            except ValueError:
                return jsonify({
                    'success': False,
                    'error': 'Expected score must be a valid number'
                }), 400
        
        # Validate file type
        allowed_extensions = {'.pdf', '.doc', '.docx', '.txt', '.jpg', '.jpeg', '.png', '.md', '.rtf', '.html', '.htm'}
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension not in allowed_extensions:
            return jsonify({
                'success': False,
                'error': f'File type {file_extension} not supported. Allowed types: {", ".join(allowed_extensions)}'
            }), 400
        
        # Create upload directory
        upload_dir = os.path.join(current_app.root_path, 'uploads', 'test_submissions')
        os.makedirs(upload_dir, exist_ok=True)
        
        # Generate unique filename
        filename = f"{current_user.id}_{int(datetime.now().timestamp())}_{file.filename}"
        file_path = os.path.join(upload_dir, filename)
        
        # Save file
        file.save(file_path)
        
        # Extract text content
        try:
            text_content = extract_text_from_file(file_path, file_extension)
            word_count = len(text_content.split()) if text_content else 0
        except Exception as e:
            logger.warning(f"Failed to extract text from test submission {filename}: {e}")
            text_content = ""
            word_count = 0
        
        # Check for duplicate content
        from src.utils.content_deduplication import check_llm_document_duplicate, get_deduplication_response, calculate_content_hash
        
        is_duplicate, existing_doc = check_llm_document_duplicate(
            user_id=current_user.id,
            content=text_content,
            document_type='test_submission',
            db_session=db.session
        )
        
        if is_duplicate:
            # Remove the uploaded file since it's a duplicate
            try:
                os.remove(file_path)
            except OSError:
                pass
            
            return jsonify(get_deduplication_response(existing_doc, "test submission")), 409
        
        # Create database record
        submission = LLMDocument(
            user_id=current_user.id,
            name=name,
            original_name=file.filename,
            stored_name=filename,
            file_type=file_extension,
            mime_type=file.mimetype or 'application/octet-stream',
            file_size=os.path.getsize(file_path),
            file_path=file_path,
            text_content=text_content,
            content_hash=calculate_content_hash(text_content),
            word_count=word_count,
            character_count=len(text_content) if text_content else 0,
            extracted_text=bool(text_content),
            type='test_submission'
        )
        
        db.session.add(submission)
        db.session.commit()
        
        logger.info(f"Test submission uploaded successfully: {name} by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': 'Test submission uploaded successfully',
            'submission': {
                'id': submission.id,
                'name': submission.name,
                'description': description,
                'expected_score': expected_score
            }
        })
        
    except Exception as e:
        logger.error(f"Error uploading test submission: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to upload test submission'
        }), 500

@llm_training_bp.route('/api/test-submissions/<submission_id>', methods=['DELETE'])
@login_required
def delete_test_submission(submission_id):
    """Delete a test submission"""
    try:
        submission = LLMDocument.query.filter_by(
            id=submission_id,
            user_id=current_user.id,
            type='test_submission'
        ).first()
        
        if not submission:
            return jsonify({
                'success': False,
                'error': 'Test submission not found'
            }), 404
        
        # Delete file
        try:
            if os.path.exists(submission.file_path):
                os.remove(submission.file_path)
        except Exception as e:
            logger.warning(f"Failed to delete test submission file {submission.file_path}: {e}")
        
        # Delete database record
        db.session.delete(submission)
        db.session.commit()
        
        logger.info(f"Test submission deleted: {submission.name} by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': 'Test submission deleted successfully'
        })
        
    except Exception as e:
        logger.error(f"Error deleting test submission: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to delete test submission'
        }), 500

@llm_training_bp.route('/api/comprehensive-reports', methods=['GET'])
@login_required
def get_comprehensive_reports():
    """Get all comprehensive reports for the current user"""
    try:
        # For now, we'll use the existing LLMTrainingReport model
        # and distinguish comprehensive reports by a special type field
        reports = LLMTrainingReport.query.filter_by(
            user_id=current_user.id
        ).filter(
            LLMTrainingReport.name.like('%Comprehensive%')
        ).order_by(LLMTrainingReport.created_at.desc()).all()
        
        reports_data = []
        for report in reports:
            reports_data.append({
                'id': report.id,
                'name': report.name,
                'description': report.description,
                'status': report.status,
                'created_at': report.created_at.isoformat(),
                'model_count': 0,  # Model count not stored in document_metadata anymore
                'file_path': report.file_path
            })
        
        return jsonify({
            'success': True,
            'reports': reports_data
        })
        
    except Exception as e:
        logger.error(f"Error getting comprehensive reports: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to load comprehensive reports'
        }), 500

@llm_training_bp.route('/api/comprehensive-reports/generate', methods=['POST'])
@login_required
def generate_comprehensive_report():
    """Generate a comprehensive report"""
    try:
        data = request.get_json()
        
        name = data.get('name', '').strip()
        description = data.get('description', '').strip()
        training_job_ids = data.get('training_job_ids', [])
        test_submission_ids = data.get('test_submission_ids', [])
        
        if not name:
            return jsonify({
                'success': False,
                'error': 'Report name is required'
            }), 400
        
        if not training_job_ids:
            return jsonify({
                'success': False,
                'error': 'At least one training job must be selected'
            }), 400
        
        if not test_submission_ids:
            return jsonify({
                'success': False,
                'error': 'At least one test submission must be selected'
            }), 400
        
        # Validate that all training jobs belong to the current user
        jobs = LLMTrainingJob.query.filter(
            LLMTrainingJob.id.in_(training_job_ids),
            LLMTrainingJob.user_id == current_user.id,
            LLMTrainingJob.status == 'completed'
        ).all()
        
        if len(jobs) != len(training_job_ids):
            return jsonify({
                'success': False,
                'error': 'Some training jobs not found or not completed'
            }), 400
        
        # Validate that all test submissions belong to the current user
        submissions = LLMDocument.query.filter(
            LLMDocument.id.in_(test_submission_ids),
            LLMDocument.user_id == current_user.id,
            LLMDocument.type == 'test_submission'
        ).all()
        
        if len(submissions) != len(test_submission_ids):
            return jsonify({
                'success': False,
                'error': 'Some test submissions not found'
            }), 400
        
        # Create comprehensive report record
        report = LLMTrainingReport(
            user_id=current_user.id,
            name=f"Comprehensive Report: {name}",
            description=description,
            format='comprehensive',
            status='generating'
            # Note: document_metadata field removed from model
        )
        
        db.session.add(report)
        db.session.commit()
        
        # Start report generation asynchronously
        from src.services.llm_training_service import LLMTrainingService
        training_service = LLMTrainingService(current_app._get_current_object())
        training_service.generate_comprehensive_report_async(report.id, jobs, submissions)
        
        logger.info(f"Comprehensive report generation started: {name} by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': 'Comprehensive report generation started',
            'report_id': report.id
        })
        
    except Exception as e:
        logger.error(f"Error generating comprehensive report: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to generate comprehensive report'
        }), 500

@llm_training_bp.route('/api/comprehensive-reports/<report_id>', methods=['DELETE'])
@login_required
def delete_comprehensive_report(report_id):
    """Delete a comprehensive report"""
    try:
        report = LLMTrainingReport.query.filter_by(
            id=report_id,
            user_id=current_user.id
        ).first()
        
        if not report:
            return jsonify({
                'success': False,
                'error': 'Comprehensive report not found'
            }), 404
        
        # Delete file if exists
        try:
            if report.file_path and os.path.exists(report.file_path):
                os.remove(report.file_path)
        except Exception as e:
            logger.warning(f"Failed to delete comprehensive report file {report.file_path}: {e}")
        
        # Delete database record
        db.session.delete(report)
        db.session.commit()
        
        logger.info(f"Comprehensive report deleted: {report.name} by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': 'Comprehensive report deleted successfully'
        })
        
    except Exception as e:
        logger.error(f"Error deleting comprehensive report: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to delete comprehensive report'
        }), 500

@llm_training_bp.route('/api/comprehensive-reports/<report_id>/download', methods=['GET'])
@login_required
def download_comprehensive_report(report_id):
    """Download a comprehensive report"""
    try:
        report = LLMTrainingReport.query.filter_by(
            id=report_id,
            user_id=current_user.id
        ).first()
        
        if not report:
            return jsonify({
                'success': False,
                'error': 'Comprehensive report not found'
            }), 404
        
        if report.status != 'completed' or not report.file_path:
            return jsonify({
                'success': False,
                'error': 'Report is not ready for download'
            }), 400
        
        if not os.path.exists(report.file_path):
            return jsonify({
                'success': False,
                'error': 'Report file not found'
            }), 404
        
        return send_file(
            report.file_path,
            as_attachment=True,
            download_name=f"{report.name}.pdf"
        )
        
    except Exception as e:
        logger.error(f"Error downloading comprehensive report: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to download comprehensive report'
        }), 500

@llm_training_bp.route('/api/test-submissions/test', methods=['POST'])
@login_required
def test_submission_with_models():
    """Test a submission with selected trained models"""
    try:
        data = request.get_json()
        
        submission_id = data.get('submission_id')
        model_ids = data.get('model_ids', [])
        
        if not submission_id:
            return jsonify({
                'success': False,
                'error': 'Submission ID is required'
            }), 400
        
        if not model_ids:
            return jsonify({
                'success': False,
                'error': 'At least one model must be selected'
            }), 400
        
        # Validate submission
        submission = LLMDocument.query.filter_by(
            id=submission_id,
            user_id=current_user.id,
            type='test_submission'
        ).first()
        
        if not submission:
            return jsonify({
                'success': False,
                'error': 'Test submission not found'
            }), 404
        
        # Validate models
        models = LLMTrainingJob.query.filter(
            LLMTrainingJob.id.in_(model_ids),
            LLMTrainingJob.user_id == current_user.id,
            LLMTrainingJob.status == 'completed'
        ).all()
        
        if len(models) != len(model_ids):
            return jsonify({
                'success': False,
                'error': 'Some models not found or not completed'
            }), 400
        
        # Start testing process asynchronously
        from src.services.llm_training_service import LLMTrainingService
        training_service = LLMTrainingService(current_app._get_current_object())
        training_service.test_submission_with_models_async(submission, models)
        
        logger.info(f"Model testing started for submission {submission.name} with {len(models)} models by user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': f'Testing started for {len(models)} models'
        })
        
    except Exception as e:
        logger.error(f"Error starting model testing: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to start model testing'
        }), 500

@llm_training_bp.route('/api/model-testing-results', methods=['GET'])
@login_required
def get_model_testing_results():
    """Get model testing results for the current user"""
    try:
        # For now, we'll create a simple structure to store testing results
        # In a real implementation, you'd have a dedicated table for this
        results = []
        
        # This is a placeholder - you'd implement actual result storage
        # For now, return empty results
        return jsonify({
            'success': True,
            'results': results
        })
        
    except Exception as e:
        logger.error(f"Error getting model testing results: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to load model testing results'
        }), 500

@llm_training_bp.route('/comprehensive-reports/<report_id>/view', methods=['GET'])
@login_required
def view_comprehensive_report(report_id):
    """View a comprehensive report in the browser"""
    try:
        report = LLMTrainingReport.query.filter_by(
            id=report_id,
            user_id=current_user.id
        ).first()
        
        if not report:
            return "Report not found", 404
        
        if report.status != 'completed':
            return "Report is not ready yet", 400
        
        # For now, return a simple HTML view
        # In a real implementation, you'd generate a comprehensive HTML report
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{report.name}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                .header {{ border-bottom: 2px solid #333; padding-bottom: 20px; margin-bottom: 30px; }}
                .section {{ margin-bottom: 30px; }}
                .model-result {{ background: #f5f5f5; padding: 15px; margin: 10px 0; border-radius: 5px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>{report.name}</h1>
                <p><strong>Description:</strong> {report.description or 'No description'}</p>
                <p><strong>Generated:</strong> {report.created_at.strftime('%Y-%m-%d %H:%M:%S')}</p>
            </div>
            
            <div class="section">
                <h2>Report Summary</h2>
                <p>This comprehensive report analyzes the performance of trained models against test submissions.</p>
                <p><strong>Models Tested:</strong> {len(training_job_ids)}</p>
                <p><strong>Test Submissions:</strong> {len(test_submission_ids)}</p>
            </div>
            
            <div class="section">
                <h2>Detailed Results</h2>
                <p><em>Detailed testing results would be displayed here in a real implementation.</em></p>
            </div>
        </body>
        </html>
        """
        
    except Exception as e:
        logger.error(f"Error viewing comprehensive report: {e}")
